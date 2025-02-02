import feedparser
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
from docx2pdf import convert  # Import docx2pdf

# RSS Feed URL
rss_url = "https://feeds.yle.fi/uutiset/v1/recent.rss?publisherIds=YLE_NEWS"

# Fetch the RSS feed
def fetch_rss_feed(url=rss_url):
    feed = feedparser.parse(url)
    return feed["entries"]

# Function to create and add a hyperlink to a cell in a DOCX table
def add_hyperlink_to_cell(cell, url, text):
    # Create the XML structure for the hyperlink
    hyperlink = OxmlElement('w:hlink')
    hyperlink.set(qn('w:anchor'), url)  # The URL
    hyperlink.set(qn('w:t'), text)  # The text to display

    # Add the hyperlink to the cell
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.clear()  # Clear any existing text
    cell_paragraph._element.append(hyperlink)

    # Set the font color to blue and underline for the link
    run = cell_paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
    run.font.underline = True

# Function to apply custom styling to the table
def style_table(table):
    # Set the table style
    table.style = 'Table Grid'
    
    # Set the header row to bold and maroon color
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Title"
    hdr_cells[1].text = "Link"
    hdr_cells[2].text = "Description"
    
    # Change the header row font color to maroon
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(128, 0, 0)  # Maroon
    
    # Set padding and extra space after descriptions
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Add space after the paragraph for description
                paragraph.paragraph_format.space_after = 12  # Adds some space after the description
                # Optional: you can add some padding around text to make it visually clearer
                cell.paragraphs[0].paragraph_format.space_before = 6  # Adds space before text

# Generate DOCX with table
def generate_docx_with_table(articles, file_name):
    doc = Document()
    doc.add_heading("YLE News Headlines", 0)
    
    # Create a table with 3 columns: Title, Link, Description
    table = doc.add_table(rows=1, cols=3)
    
    # Apply custom styling
    style_table(table)
    
    # Loop through articles and add them to the table
    for article in articles:
        title = article.get("title", "No title")
        link = article.get("link", "#")
        description = article.get("description", "No description available")
        
        # Add the title to the table
        row_cells = table.add_row().cells
        row_cells[0].text = title
        
        # Add the link as a clickable hyperlink
        add_hyperlink_to_cell(row_cells[1], link, link)
        
        # Add the description
        row_cells[2].text = description
        
        # Add an empty row to separate the sections
        table.add_row()  # This will create an empty row
    
    # Save DOCX to the specified path
    doc.save(file_name)

# Convert DOCX to PDF using docx2pdf
def convert_docx_to_pdf(docx_file, output_pdf_file):
    try:
        # Use docx2pdf to convert DOCX to PDF
        convert(docx_file, output_pdf_file)
        print(f"PDF created successfully: {output_pdf_file}")
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")

# Main function
def main():
    articles = fetch_rss_feed()
    
    # Get today's date and format it as DD_MM_YYYY
    today = datetime.now().strftime("%d_%m_%Y")
    
    # Specify the drive location and generate file name
    drive_folder = "F:/SELF/YLE/Docs"  # Update this to your folder path
    file_name = f"YLE_News_Headlines_{today}.docx"
    
    # Ensure the directory exists
    if not os.path.exists(drive_folder):
        os.makedirs(drive_folder)
    
    # Full path to save the DOCX file
    docx_file_path = os.path.join(drive_folder, file_name)
    
    # Generate DOCX with table
    generate_docx_with_table(articles, docx_file_path)
    
    print(f"DOCX generated and saved to: {docx_file_path}")
    
    # Specify the PDF file path
    pdf_file_path = docx_file_path.replace(".docx", ".pdf")
    
    # Convert DOCX to PDF
    convert_docx_to_pdf(docx_file_path, pdf_file_path)
    print(f"PDF saved to: {pdf_file_path}")

if __name__ == "__main__":
    main()